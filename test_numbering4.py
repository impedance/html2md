#!/usr/bin/env python3
"""Find all potential headings and their context."""

import sys
from docx import Document
import re


def find_all_numbered_content(docx_path: str):
    """Find all numbered content."""
    doc = Document(docx_path)
    
    print(f"Finding numbered content in {docx_path}...")
    print("=" * 60)
    
    found_content = False
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        
        # Look for specific patterns that appear in your screenshot
        patterns = [
            r'^\d+\.\d+\.\d+\.\d+\s+[А-Я]',  # Like "4.1.2.1 Подготовка"
            r'^\d+\.\d+\.\d+\s+[А-Я]',       # Like "4.1.2 Something"
            r'^\d+\.\d+\s+[А-Я]',            # Like "4.1 Something"
            r'^\d+\s+[А-Я][а-я]',            # Like "4 Настройка"
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                found_content = True
                number_part = re.match(r'^(\d+(?:\.\d+)*)', text).group(1)
                level = len(number_part.split('.'))
                
                print(f"Para {i:4d}: Level H{level} Style='{paragraph.style.name}'")
                print(f"         Number: '{number_part}'")
                print(f"         Text: '{text}'")
                print()
                break
    
    if not found_content:
        print("Looking for any text containing '4.1.2.1' or similar patterns...")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if '4.1.2.1' in text or '4.1.2' in text or 'Подготовка' in text:
                print(f"Para {i:4d}: Style='{paragraph.style.name}'")
                print(f"         Text: '{text[:100]}{'...' if len(text) > 100 else ''}'")
                print()


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python test_numbering4.py <docx_file>")
        sys.exit(1)
    
    find_all_numbered_content(sys.argv[1])