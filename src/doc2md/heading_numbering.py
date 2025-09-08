"""Extract heading numbering information from DOCX files."""

from __future__ import annotations

import re
from typing import Dict, List, Tuple
from docx import Document


def extract_heading_numbering_from_toc(docx_path: str) -> Dict[str, str]:
    """
    Extract heading numbering from the Table of Contents in a DOCX file.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Dictionary mapping heading text (without numbers) to their numbers
        Example: {"Общие сведения": "1", "Назначение": "1.1", "Подготовка конфигурационных файлов": "4.1.2.1"}
    """
    doc = Document(docx_path)
    numbering_map: Dict[str, str] = {}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = paragraph.style
        if style is None or not style.name.startswith("toc"):
            continue

            # Check if this is a TOC entry with numbering
            # Pattern to match numbered TOC entries like "4.1.2.1 Подготовка конфигурационных файлов\t42"
            match = re.match(r"^(\d+(?:\.\d+)*)\s+([^\t]+)(?:\t\d+)?$", text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                numbering_map[title] = number

    return numbering_map


def extract_heading_structure_from_toc(docx_path: str) -> List[Tuple[int, str, str]]:
    """
    Extract complete heading structure from TOC.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        List of tuples (level, number, title) sorted by document order
        Example: [(1, "1", "Общие сведения"), (2, "1.1", "Назначение"), (4, "4.1.2.1", "Подготовка")]
    """
    doc = Document(docx_path)
    headings = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = paragraph.style
        if style is None or not style.name.startswith("toc"):
            continue

        # Extract level from style name (toc 1, toc 2, etc.)
        level_match = re.search(r"toc (\d+)", style.name)
        if not level_match:
            continue

        level = int(level_match.group(1))

        # Extract number and title
        match = re.match(r"^(\d+(?:\.\d+)*)\s+([^\t]+)(?:\t\d+)?$", text)
        if match:
            number = match.group(1)
            title = match.group(2).strip()
            headings.append((level, number, title))

    return headings


def get_heading_number_for_text(text: str, numbering_map: Dict[str, str]) -> str | None:
    """
    Find the heading number for a given text by fuzzy matching against the numbering map.

    Args:
        text: The heading text to match
        numbering_map: Dictionary from extract_heading_numbering_from_toc

    Returns:
        The heading number if found, None otherwise
    """
    text = text.strip()

    # First try exact match
    if text in numbering_map:
        return numbering_map[text]

    # Try fuzzy matching - look for text that contains the same words
    text_words = set(text.lower().split())
    best_match = None
    best_score = 0.0

    for toc_title, number in numbering_map.items():
        toc_words = set(toc_title.lower().split())

        # Calculate similarity score (intersection over union)
        intersection = len(text_words & toc_words)
        union = len(text_words | toc_words)

        if union > 0:
            score = intersection / union
            if score > best_score and score > 0.5:  # At least 50% similarity
                best_score = score
                best_match = number

    return best_match


def add_numbering_to_html(html_content: str, docx_path: str) -> str:
    """
    Add heading numbering to HTML content based on DOCX TOC.
    Uses regex to find and replace __RefHeading patterns with proper heading tags.

    Args:
        html_content: HTML content from Mammoth conversion
        docx_path: Path to original DOCX file

    Returns:
        HTML content with numbered headings
    """
    import re

    # Extract heading structure from DOCX
    heading_structure = extract_heading_structure_from_toc(docx_path)

    if not heading_structure:
        return html_content  # No headings found

    result = html_content

    # Replace each anchor + title pair with a proper heading tag
    for level, number, title in heading_structure:
        escaped_title = re.escape(title)
        pattern = re.compile(
            rf'<a id="__RefHeading___\d+"></a>\s*(?:<[^>]+>\s*)*{escaped_title}',
            flags=re.IGNORECASE,
        )
        replacement = f"<h{level}>{number} {title}</h{level}>"
        result, count = pattern.subn(replacement, result, count=1)

        if count == 0:
            # Remove unmatched anchor to avoid leaking into output
            result = re.sub(r'<a id="__RefHeading___\d+"></a>', "", result, count=1)

    # Clean up any remaining reference anchors
    result = re.sub(r'<a id="__RefHeading___\d+"></a>', "", result)
    return result


if __name__ == "__main__":
    # Test the functionality
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        print("Extracting heading numbering from TOC...")
        numbering_map = extract_heading_numbering_from_toc(docx_path)

        print("Found numbering:")
        for title, number in sorted(numbering_map.items(), key=lambda x: x[1]):
            print(f"  {number}: {title}")

        print(f"\nTotal headings: {len(numbering_map)}")

        print("\nHeading structure:")
        structure = extract_heading_structure_from_toc(docx_path)
        for level, number, title in structure[:20]:  # Show first 20
            indent = "  " * (level - 1)
            print(f"{indent}H{level} {number} {title}")
